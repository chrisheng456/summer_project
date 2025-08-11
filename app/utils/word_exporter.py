import json
from typing import Union
from datetime import datetime
from docx import Document
from docx.shared import Pt


def format_time(timestr):
    if not timestr:
        return ""
    try:
        if timestr.endswith("Z"):
            timestr = timestr[:-1]
        dt = datetime.fromisoformat(timestr)
        return dt.strftime("%Y-%m-%d %H:%M")
    except Exception:
        return timestr


def export_to_word(result_json: Union[str, dict]) -> bytes:
    if isinstance(result_json, str):
        data = json.loads(result_json)
    else:
        data = result_json
    doc = Document()
    doc.add_heading(data.get("name", "Meeting Minutes"), 0)
    # 会议基本信息
    for k in ["date", "startTime", "location"]:
        v = data.get(k, "")
        if k in ("date", "startTime"):
            v = format_time(v)
        if v:
            p = doc.add_paragraph()
            p.add_run(f"{k.title()}: ").bold = True
            p.add_run(str(v))
    # 与会者
    attendees = data.get("attendees", [])
    if attendees:
        doc.add_heading("Attendees", level=1)
        for att in attendees:
            doc.add_paragraph(f"- {att.get('name', '')}", style="List Bullet")
    # 议程
    agenda = data.get("agenda", [])
    for idx, item in enumerate(agenda, 1):
        doc.add_heading(
            f"Agenda {item.get('number', idx)}: {item.get('title', '')}",
            level=2,
        )
        if item.get("owner"):
            doc.add_paragraph(f"Owner: {item['owner']}")
        if item.get("calculatedStartTime"):
            doc.add_paragraph(
                f"Start: {format_time(item['calculatedStartTime'])}"
            )
        if item.get("lengthMinutes"):
            doc.add_paragraph(f"Duration: {item['lengthMinutes']} min")
        if item.get("label"):
            doc.add_paragraph(
                f"Label: {item['label']} (score: {item.get('label_score', '')})"
            )
        if item.get("explanation"):
            p = doc.add_paragraph("Explanation: ")
            p.add_run(item["explanation"]).italic = True
        if item.get("summary"):
            p = doc.add_paragraph("Summary: ")
            p.add_run(item["summary"]).bold = True
        # lines
        lines = item.get("lines", [])
        for ln in lines:
            speaker = ln.get("speaker", "")
            start = ln.get("start", "")
            end = ln.get("end", "")
            text = ln.get("text", "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run(f"{speaker} [{start:.2f}-{end:.2f}s]: ").italic = True
            p.add_run(text)
    # 导出为bytes
    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
