import azure.cognitiveservices.speech as speechsdk
import json
import uuid
from docx import Document
import threading

# 配置 Azure 语音服务（去除尖括号，替换为你的实际值）
speech_key = "3WiTSlcuLgPpMIoMfsJnNvm0exinJD1NxoOobVQzEXTc7OenGnZhJQQJ99BFAClhwhEXJ3w3AAAYACOGBTOW"
service_region = "ukwest"  # 如 "eastus"
speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)

# 音频文件路径
audio_path = "test.wav"  # 替换为你的音频文件路径
audio_input = speechsdk.AudioConfig(filename=audio_path)
recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_input)

# 用于存储识别结果
full_text_lines = []
# 用于等待识别结束
stop_event = threading.Event()

# 识别到语音时的回调
def recognized_handler(evt):
    if evt.result.reason == speechsdk.ResultReason.RecognizedSpeech:
        print(f"RECOGNIZED: {evt.result.text}")
        full_text_lines.append(evt.result.text)
    elif evt.result.reason == speechsdk.ResultReason.NoMatch:
        print("NOMATCH: 无法识别语音片段。")

# 识别取消时的回调
def canceled_handler(evt):
    cancellation = speechsdk.CancellationDetails(evt)
    print(f"CANCELED: 原因={cancellation.reason}, 详情={cancellation.error_details}")
    stop_event.set()

# 会话结束时的回调
def session_stopped_handler(evt):
    print("SESSION STOPPED: 识别已完成。")
    stop_event.set()

# 订阅事件
recognizer.recognized.connect(recognized_handler)
recognizer.canceled.connect(canceled_handler)
recognizer.session_stopped.connect(session_stopped_handler)

# 启动连续识别
print(f"开始连续识别音频：{audio_path}…")
recognizer.start_continuous_recognition()

# 等待会话结束事件
stop_event.wait()

# 停止连续识别
recognizer.stop_continuous_recognition()

# 合并识别文本
text = "\n".join(full_text_lines)
print("最终转写文本：")
print(text)

# 保存为 Word 文档
docx_filename = "meeting_transcription.docx"
doc = Document()
doc.add_heading("会议转录结果", level=1)
for line in text.split("\n"):
    doc.add_paragraph(line)
doc.save(docx_filename)
print(f"已保存 Word 文档：{docx_filename}")

# 构造 JSON 输出结构
output = {
    "id": str(uuid.uuid4()),
    "transcription": text,
    "abstract_summary": None,
    "key_points": [],
    "action_items": [],
    "sentiment": None,
    "attachment": []
}

# 保存为 JSON 文件
json_filename = "meeting_minutes.json"
with open(json_filename, "w", encoding="utf-8") as f:
    json.dump(output, f, ensure_ascii=False, indent=2)
print(f"已保存 JSON 文件：{json_filename}")
