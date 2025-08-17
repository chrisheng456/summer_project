from __future__ import annotations
from typing import Any, Dict, List, Tuple, Optional
import io, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

_ZWS = "\u200b"

def _soft_wrap(text: str, every: int = 24) -> str:
    def wrap_token(tok: str) -> str:
        if len(tok) <= every or any(ch.isspace() for ch in tok): return tok
        return _ZWS.join(tok[i:i+every] for i in range(0, len(tok), every))
    text = (text or "").replace("\u00A0"," ").replace("\u2009"," ").replace("\u202F"," ").replace("\u2060","")
    parts = re.findall(r"\S+|\s+", text, flags=re.UNICODE)
    return "".join(wrap_token(p) if not p.isspace() else p for p in parts)

def _norm(s: Optional[str]) -> str:
    if not s: return ""
    s = s.replace("\r\n","\n").replace("\r","\n")
    s = re.sub(r"[^\S\n\t]+"," ", s)
    s = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]","", s)
    return _soft_wrap(s).strip()

def _fmt_time(v: Any) -> str:
    if v is None: return ""
    try:
        x = float(v); m = int(x//60); s = int(round(x-m*60))
        return f"{m:02d}:{s:02d}"
    except Exception:
        return _norm(str(v))

def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(_norm(text), level=level); p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_kv_table(doc: Document, rows: List[Tuple[str,str]]):
    table = doc.add_table(rows=1, cols=2); table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells; hdr[0].text = "Field"; hdr[1].text = "Value"
    for k,v in rows:
        r = table.add_row().cells; r[0].text = _norm(k); r[1].text = _norm(v)

def _add_lines_table(doc: Document, lines: List[Dict[str,Any]]):
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 2"
    hdr = tbl.rows[0].cells; hdr[0].text="Start"; hdr[1].text="End"; hdr[2].text="Speaker"; hdr[3].text="Text"
    for ln in lines:
        r = tbl.add_row().cells
        r[0].text = _fmt_time(ln.get("start")); r[1].text = _fmt_time(ln.get("end"))
        r[2].text = _norm(ln.get("speaker","")); r[3].text = _norm(ln.get("text",""))

def build_meeting_docx(meeting_detail: Dict[str,Any]) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    # 顶部
    _add_heading(doc, meeting_detail.get("name","Meeting"), level=1)
    rows: List[Tuple[str,str]] = [
        ("Meeting ID", str(meeting_detail.get("id",""))),
        ("Date", _norm(meeting_detail.get("date") or "")),
        ("Start Time", _norm(meeting_detail.get("startTime") or "")),
        ("Location", _norm(meeting_detail.get("location") or "")),
    ]
    attendees = meeting_detail.get("attendees") or []
    if attendees:
        rows.append(("Attendees", ", ".join(_norm(x.get("name","")) for x in attendees)))
    _add_kv_table(doc, rows)

    agenda: List[Dict[str,Any]] = meeting_detail.get("agenda") or []
    if not agenda:
        doc.add_paragraph("No agenda items.")

    # 逐个 agenda
    for it in agenda:
        number = str(it.get("number","")); title = _norm(it.get("title",""))
        _add_heading(doc, f"Agenda {number}: {title}", level=2)

        meta: List[Tuple[str,str]] = []
        if it.get("owner"): meta.append(("Owner", _norm(it.get("owner"))))
        if it.get("calculatedStartTime"): meta.append(("Start", _norm(it.get("calculatedStartTime"))))
        if it.get("lengthMinutes") is not None: meta.append(("Duration", f"{it.get('lengthMinutes')} min"))
        if it.get("label"): meta.append(("Label", f"{_norm(it.get('label'))} (score: {it.get('label_score','')})"))
        if meta: _add_kv_table(doc, meta)

        if it.get("explanation"):
            _add_heading(doc, "Explanation", level=3)
            doc.add_paragraph(_norm(it.get("explanation")))
        if it.get("summary"):
            _add_heading(doc, "Summary", level=3)
            p = doc.add_paragraph(_norm(it.get("summary"))); p.paragraph_format.space_after = Pt(6)

        lines = it.get("lines") or []
        if lines:
            _add_heading(doc, "Lines", level=3)
            _add_lines_table(doc, lines)
        else:
            doc.add_paragraph("No lines.")

    out = io.BytesIO(); doc.save(out); return out.getvalue()

def export_docx_from_payload(payload: Dict[str,Any]) -> bytes:
    if not payload: raise ValueError("empty payload")
    md = payload.get("customer_meeting_detail") or payload.get("meeting_detail") or payload
    if not isinstance(md, dict) or not md.get("agenda"):
        raise ValueError("not a valid meeting detail JSON (missing 'agenda')")
    return build_meeting_docx(md)

# 兼容旧名
export_to_word = export_docx_from_payload
