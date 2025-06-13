"""
speech_to_text.py
------------------------------------
• TrueText（自动标点 + 首字母大写）
• 说话人分离（Diarization）
• 保存为 Word 和 JSON（带时间戳文件名）
"""

import os
import json
import uuid
import threading
from datetime import datetime
from pathlib import Path

import azure.cognitiveservices.speech as speechsdk     # pip install -U azure-cognitiveservices-speech
from docx import Document                              # pip install python-docx

# ─────────────────────────────────────────────────────
# 1. 读取密钥：系统环境变量，或取消注释使用 .env
# ─────────────────────────────────────────────────────
# from dotenv import load_dotenv
# load_dotenv()                                       # 若项目根有 .env，可启用

speech_key     = os.getenv("AZURE_SPEECH_KEY")
service_region = os.getenv("AZURE_SPEECH_REGION", "eastus")

if not speech_key:
    raise RuntimeError("❌ 找不到 AZURE_SPEECH_KEY，请先在系统变量或 .env 中设置")

# ─────────────────────────────────────────────────────
# 2. SpeechConfig & 功能开关
# ─────────────────────────────────────────────────────
speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)

# TrueText：自动标点 / 大小写
speech_config.set_property(
    speechsdk.PropertyId.SpeechServiceResponse_PostProcessingOption,
    "TrueText"                   # 或 "TrueText-NoFiller"
)

# 说话人分离
speech_config.set_property(
    speechsdk.PropertyId.SpeechServiceResponse_DiarizeIntermediateResults,
    "true"
)

speech_config.request_word_level_timestamps()          # 可选：字级时间戳

# ─────────────────────────────────────────────────────
# 3. 音频文件
# ─────────────────────────────────────────────────────
AUDIO_PATH = Path("test.wav")      # ← 改成自己的音频文件
if not AUDIO_PATH.exists():
    raise FileNotFoundError(f"找不到音频文件：{AUDIO_PATH.resolve()}")

audio_config = speechsdk.AudioConfig(filename=str(AUDIO_PATH))

# ConversationTranscriber 支持 Diarization
transcriber = speechsdk.transcription.ConversationTranscriber(
    speech_config=speech_config,
    audio_config=audio_config
)

# ─────────────────────────────────────────────────────
# 4. 事件回调
# ─────────────────────────────────────────────────────
lines = []                          # 保存每一句
done  = threading.Event()           # 等待识别结束用

def _on_transcribed(evt: speechsdk.SpeechRecognitionEventArgs):
    if evt.result.reason == speechsdk.ResultReason.RecognizedSpeech:
        lines.append({
            "speaker" : evt.result.speaker_id or "Unknown",
            "text"    : evt.result.text,
            "offset"  : evt.result.offset,      # 100-ns
            "duration": evt.result.duration,
        })
        print(f"[{lines[-1]['speaker']}] {lines[-1]['text']}")
    elif evt.result.reason == speechsdk.ResultReason.NoMatch:
        print("NOMATCH: 该片段无法识别")

def _on_session_stopped(_):
    print("=== 识别结束 ===")
    done.set()

def _on_canceled(evt):
    details = speechsdk.CancellationDetails(evt)
    print(f"CANCELED: {details.reason} / {details.error_details}")
    done.set()

transcriber.transcribed.connect(_on_transcribed)
transcriber.session_stopped.connect(_on_session_stopped)
transcriber.canceled.connect(_on_canceled)

# ─────────────────────────────────────────────────────
# 5. 开始转写 & 等待结束
# ─────────────────────────────────────────────────────
print(f"▶ 开始识别 {AUDIO_PATH} ...")
transcriber.start_transcribing_async()
done.wait()
transcriber.stop_transcribing_async()

# ─────────────────────────────────────────────────────
# 6. 整理 & 生成带时间戳文件名
# ─────────────────────────────────────────────────────
lines.sort(key=lambda x: x["offset"])
plain_text = "\n".join(f"{ln['speaker']}: {ln['text']}" for ln in lines)

stamp      = datetime.now().strftime("%Y%m%d_%H%M%S")
docx_name  = f"meeting_transcription_{stamp}.docx"
json_name  = f"meeting_minutes_{stamp}.json"
out_id     = str(uuid.uuid4())

# (a) Word
doc = Document()
doc.add_heading("会议逐字稿", level=1)
for ln in lines:
    doc.add_paragraph(f"{ln['speaker']}: {ln['text']}")
doc.save(docx_name)
print(f"✅ 已保存 Word：{docx_name}")

# (b) JSON
output = {
    "id"              : out_id,
    "transcription"   : plain_text,
    "lines"           : lines,
    "abstract_summary": None,
    "key_points"      : [],
    "action_items"    : [],
    "sentiment"       : None,
    "attachment"      : [],
}
with open(json_name, "w", encoding="utf-8") as f:
    json.dump(output, f, ensure_ascii=False, indent=2)
print(f"✅ 已保存 JSON：{json_name}")
