"""
• TrueText (automatic punctuation + capitalization)
• Speaker diarization
• Save as both Word and JSON (with start/end/text only)
• Automatically convert .m4a to 16kHz mono WAV via ffmpeg
"""
import os
import json
import threading
import subprocess
from datetime import datetime
from pathlib import Path

import azure.cognitiveservices.speech as speechsdk  # pip install -U azure-cognitiveservices-speech
from docx import Document                         # pip install python-docx

# 1. Read Azure keys from environment
speech_key     = os.getenv("AZURE_SPEECH_KEY")
service_region = os.getenv("AZURE_SPEECH_REGION", "ukwest")
if not speech_key:
    raise RuntimeError(" AZURE_SPEECH_KEY not found. Please set it in environment variables or .env file.")

# 2. Configure Azure Speech settings
speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)
speech_config.set_property(
    speechsdk.PropertyId.SpeechServiceResponse_PostProcessingOption,
    "TrueText"
)
speech_config.set_property(
    speechsdk.PropertyId.SpeechServiceResponse_DiarizeIntermediateResults,
    "true"
)
speech_config.request_word_level_timestamps()

# 3. Convert input audio (.m4a) into 16kHz mono WAV
BASE_DIR = Path(__file__).parent
src_path = BASE_DIR / "Trustee Meeting Recording (30 June 2025) V1.m4a"
wav_path = BASE_DIR / "tmp_converted.wav"

if not src_path.exists():
    raise FileNotFoundError(f"Source audio file not found: {src_path.resolve()}")

print(f" Converting audio: {src_path.name} → {wav_path.name} (16kHz mono WAV)")
subprocess.run([
    "ffmpeg", "-y",
    "-i", str(src_path),
    "-ac", "1",
    "-ar", "16000",
    str(wav_path)
], check=True)

# 4. Prepare Azure transcriber
audio_config = speechsdk.AudioConfig(filename=str(wav_path))
transcriber  = speechsdk.transcription.ConversationTranscriber(
    speech_config=speech_config,
    audio_config=audio_config
)

lines   = []             # store recognized lines with timestamps
is_done = threading.Event()

def _on_transcribed(evt: speechsdk.SpeechRecognitionEventArgs):
    """Callback for recognized speech."""
    if evt.result.reason != speechsdk.ResultReason.RecognizedSpeech:
        return
    text = evt.result.text.strip()
    if not text:
        return
    start_sec = evt.result.offset / 10_000_000
    end_sec   = start_sec + (evt.result.duration / 10_000_000)
    lines.append({"start": start_sec, "end": end_sec, "text": text})
    print(f"[{start_sec:.2f}s - {end_sec:.2f}s] {text}")

def _on_session_stopped(_):
    print("=== Transcription complete ===")
    is_done.set()

def _on_canceled(evt):
    details = speechsdk.CancellationDetails(evt)
    print(f"CANCELED: {details.reason} / {details.error_details}")
    is_done.set()

transcriber.transcribed.connect(_on_transcribed)
transcriber.session_stopped.connect(_on_session_stopped)
transcriber.canceled.connect(_on_canceled)

# 5. Start transcription and wait for completion
print(f"▶ Starting transcription for {wav_path.name} ...")
transcriber.start_transcribing_async()
is_done.wait()
transcriber.stop_transcribing_async()

# 6. Save results (Word + JSON) using base filename
lines.sort(key=lambda x: x["start"])
base_name = src_path.stem

docx_name = f"{base_name}.docx"
json_name = f"{base_name}.json"

# (a) Save as Word
doc = Document()
doc.add_heading("Meeting Transcript", level=1)
for ln in lines:
    doc.add_paragraph(f"[{ln['start']:.2f}s - {ln['end']:.2f}s] {ln['text']}")
doc.save(docx_name)
print(f" Word file saved: {docx_name}")

# (b) Save as JSON
with open(json_name, "w", encoding="utf-8") as f:
    json.dump({"lines": lines}, f, ensure_ascii=False, indent=2)
print(f" JSON file saved: {json_name}")
