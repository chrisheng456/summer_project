#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import re
import argparse
from pathlib import Path
from datetime import timedelta
from collections import defaultdict
from docx import Document

def load_json(path: Path):
    """从文件加载 JSON"""
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def compute_speaker_times(diarized_records):
    """
    统计每个说话人的最早发言（排除 start==0 的假段落）和总时长
    返回：{ speaker: {"start_time": float, "duration": float}, ... }
    """
    segs_by_spk = defaultdict(list)
    for r in diarized_records:
        spk = r.get("speaker", "Unknown")
        s, e = r.get("start", 0.0), r.get("end", 0.0)
        segs_by_spk[spk].append((s, e - s))
    out = {}
    for spk, segs in segs_by_spk.items():
        # 先过滤掉 start==0 的“假”分段
        valid = [(s, d) for s, d in segs if s > 0]
        if valid:
            starts, durs = zip(*valid)
        else:
            # 如果该人所有分段都从 0 开始，就用原始数据
            starts, durs = zip(*segs)
        out[spk] = {
            "start_time": min(starts),
            "duration":   sum(durs)
        }
    return out

def fmt_time(sec: float) -> str:
    """把秒数转成 hh:mm:ss 格式"""
    return str(timedelta(seconds=int(sec)))

def split_summary(combined: str):
    """
    把形如
      "see ...; decided: \"...\"; raised concern: \"...\""
    的 summary 字符串，拆成三部分：
      action_summary, decision_summary, concern_summary
    """
    # 匹配 decided 和 raised concern
    dec_m = re.search(r'; decided: "([^"]*)"', combined)
    con_m = re.search(r'; raised concern: "([^"]*)"', combined)

    # action_summary 部分是从开头到第一个 '; decided' 或 '; raised concern'
    end_idx = min(
        dec_m.start() if dec_m else len(combined),
        con_m.start() if con_m else len(combined)
    )
    action_summary = combined[:end_idx].strip().rstrip(';')

    decision_summary = dec_m.group(1) if dec_m else ""
    concern_summary  = con_m.group(1) if con_m else ""
    return action_summary, decision_summary, concern_summary

def build_agenda_by_speaker(summaries, time_info):
    """
    按说话人分段生成议程结构：
    [
      {
        "speaker": "...",
        "start_time": "hh:mm:ss",
        "duration":   "hh:mm:ss",
        "items": [
           {
             "id": 1,
             "action_summary": "...",
             "decision_summary": "...",
             "concern_summary": "...",
             "actions": [...],
             "decisions": [...],
             "conflicts": [...]
           },
           ...
        ]
      },
      ...
    ]
    """
    grp = defaultdict(list)
    for x in summaries:
        grp[x["speaker"]].append(x)

    sections = []
    for spk, recs in grp.items():
        ti = time_info.get(spk, {"start_time": 0, "duration": 0})
        sec = {
            "speaker":     spk,
            "start_time":  fmt_time(ti["start_time"]),
            "duration":    fmt_time(ti["duration"]),
            "items":       []
        }
        for idx, x in enumerate(recs, start=1):
            sec["items"].append({
                "id": idx,
                "actions": x.get("actions", []),
                "decisions": x.get("decisions", []),
                "conflicts": x.get("conflicts", [])
            })

        sections.append(sec)
    return sections

def save_agenda_json(sections, out_path: Path):
    """
    将按说话人分段的议程写入 JSON，
    并且在每个分组开头加上 "section" 序号
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)

    numbered = []
    for idx, sec in enumerate(sections, start=1):
        # 先把 section 放前面，再展开原 sec 的所有字段
        sec_numbered = {"section": idx, **sec}
        numbered.append(sec_numbered)

    output = {"agenda_by_speaker": numbered}
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)



def save_agenda_docx(sections, out_path: Path):
    """
    将按说话人分段的议程写入 DOCX，
    只展示 actions / decisions / conflicts，不再引用已删除的 summary 字段
    """
    from docx import Document

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("会议议程（按说话人分段）", level=1)

    for sec_idx, sec in enumerate(sections, start=1):
        # 二级标题：Section N. Speaker + 时间信息
        h = doc.add_heading(level=2)
        h.add_run(f"Section {sec_idx}. ").bold = True
        h.add_run(f"{sec['speaker']} ").bold = True
        h.add_run(f"[起始：{sec['start_time']}，总时长：{sec['duration']}]")

        # 列出该说话人所有议题，输出编号 + 说话人
        for it in sec["items"]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{it['id']}. {sec['speaker']}").bold = True

            # 列出所有 actions
            for act in it.get("actions", []):
                pa = doc.add_paragraph(style="List Bullet")
                pa.add_run("Action: " + act)

            # 列出所有 decisions
            for dec in it.get("decisions", []):
                pd = doc.add_paragraph(style="List Bullet")
                pd.add_run("Decision: " + dec)

            # 列出所有 conflicts
            for cf in it.get("conflicts", []):
                pc = doc.add_paragraph(style="List Bullet")
                pc.add_run("Conflict: " + cf)

    doc.save(str(out_path))



if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="按说话人分段生成会议议程")
    this_dir = Path(__file__).parent

    parser.add_argument(
        "--diarized",
        help="说话人分段 JSON 文件路径（含 start/end 字段）",
        default=str(this_dir / "dataset" / "ICSI" / "diarized_json" / "Bed002_diarized.json")
    )
    parser.add_argument(
        "--summary",
        help="summary_json.py 生成的 summary 文件路径",
        default=str(this_dir / "dataset" / "ICSI" / "after_json" / "Bed002_summary.json")
    )
    parser.add_argument(
        "--out_json",
        help="输出的 Agenda JSON 文件路径",
        default=str(this_dir / "output" / "agenda.json")
    )
    parser.add_argument(
        "--out_docx",
        help="输出的 Agenda DOCX 文件路径",
        default=str(this_dir / "output" / "agenda.docx")
    )
    args = parser.parse_args()

    # 确保输出目录存在
    Path(args.out_json).parent.mkdir(parents=True, exist_ok=True)
    Path(args.out_docx).parent.mkdir(parents=True, exist_ok=True)

    diarized = load_json(Path(args.diarized))
    summaries = load_json(Path(args.summary))
    time_info = compute_speaker_times(diarized)
    sections = build_agenda_by_speaker(summaries, time_info)

    save_agenda_json(sections, Path(args.out_json))
    save_agenda_docx(sections,  Path(args.out_docx))

    print("✅ 按说话人分段的议程已生成：")
    print("   JSON →", args.out_json)
    print("   DOCX →", args.out_docx)
