from fastapi import FastAPI, UploadFile, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from transformers import pipeline
from docx import Document
from fpdf import FPDF
from pathlib import Path
import shutil, os, json, datetime, torch, mimetypes, whisper
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

# ==== 引入语义搜索组件 ====
from evaluate.qdrant_search import HybridSearcher, COLLECTION

# 初始化 FastAPI
app = FastAPI(title="会议纪要助手（支持语义搜索）")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 文件存储目录
TEMP_DIR = "temp/temp"
Path(TEMP_DIR).mkdir(exist_ok=True)
Path("temp/output").mkdir(exist_ok=True)

# Whisper模型
model_whisper = whisper.load_model("base")

# NLP管道
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6",
                      device=0 if torch.cuda.is_available() else -1)
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli",
                      device=0 if torch.cuda.is_available() else -1)

# ==== 语义搜索模型初始化 ====
hs = HybridSearcher()


# ==== 数据模型定义 ====
class TranscriptItem(BaseModel):
    speaker: str
    text: str
    time: str = None


class TranscriptData(BaseModel):
    transcript: list[TranscriptItem]


class SearchRequest(BaseModel):
    query: str


# ---------- 1. 上传音频 ----------
@app.post("/upload_audio")
async def upload_audio(file: UploadFile):
    safe_name = file.filename.replace(" ", "_").replace("(", "").replace(")", "").replace("@", "")
    save_path = os.path.join(TEMP_DIR, safe_name)
    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    return {"message": "文件接收成功", "filename": safe_name}
    print(f"✅ 文件实际保存为: {save_path}")


# ---------- 2. 自动转录 ----------
@app.post("/transcribe")
async def transcribe_audio(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)

    print("🛠 正在尝试读取文件路径：", file_path)
    if not os.path.exists(file_path):
        return {"error": f"[WinError 2] 系统找不到指定的文件：{file_path}"}

    try:
        result = model_whisper.transcribe(file_path)
        return {"transcription": result["text"]}
    except Exception as e:
        return {"error": str(e)}



# ---------- 3. 自动生成纪要 ----------
@app.post("/generate_minutes")
async def generate_minutes(data: TranscriptData):
    utterances = data.transcript
    all_text = "\n".join([utt.text for utt in utterances if utt.text.strip()])
    if not all_text:
        return {"error": "Empty transcript"}

    summary_out = summarizer(all_text, max_length=150, min_length=40, do_sample=False)
    summary = summary_out[0]["summary_text"]

    agenda_items, action_items, decisions, conflicts = [], [], [], []
    for utt in utterances:
        result = classifier(utt.text, candidate_labels=["agenda", "action", "decision", "conflict of interest"])
        top_label = result["labels"][0]
        score = result["scores"][0]
        if top_label == "agenda" and score > 0.6:
            agenda_items.append(utt.text)
        elif top_label == "action" and score > 0.6:
            action_items.append(utt.text)
        elif top_label == "decision" and score > 0.6:
            decisions.append(utt.text)
        elif top_label == "conflict of interest" and score > 0.6:
            conflicts.append(utt.text)

    result = {
        "summary": summary,
        "agenda_items": agenda_items,
        "action": action_items,
        "decision": decisions,
        "conflict_of_interest": conflicts,
        "speaker": [utt.speaker for utt in utterances],
        "time": [utt.time for utt in utterances if utt.time],
        "transcription": [utt.text for utt in utterances]
    }

    with open("output/generated_minutes.json", "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return result


# ---------- 4. 导出 Markdown ----------
@app.post("/export_markdown")
async def export_markdown(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.md"
    content = f"""# 📝 Meeting Summary

## 🟨 Summary  
{minutes_data.get('summary', '')}

## 🔹 Agenda Items  
""" + "\n".join(f"- {item}" for item in minutes_data.get("agenda_items", [])) + """

## ✅ Actions  
""" + "\n".join(f"- [ ] {a}" for a in minutes_data.get("action", [])) + """

## 📌 Decisions  
""" + "\n".join(f"- {d}" for d in minutes_data.get("decision", [])) + """

## ⚠️ Conflict of Interest  
""" + "\n".join(f"- {c}" for c in minutes_data.get("conflict_of_interest", []))

    path = os.path.join(TEMP_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return {"message": "Markdown exported", "filename": filename}


# ---------- 5. 导出 Word ----------
@app.post("/export_word")
async def export_word(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.docx"
    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(minutes_data.get("summary", ""))
    doc.add_heading("Agenda Items", level=2)
    for item in minutes_data.get("agenda_items", []):
        doc.add_paragraph(f"• {item}", style='List Bullet')
    doc.add_heading("Actions", level=2)
    for a in minutes_data.get("action", []):
        doc.add_paragraph(f"☐ {a}", style='List Bullet')
    doc.add_heading("Decisions", level=2)
    for d in minutes_data.get("decision", []):
        doc.add_paragraph(f"• {d}", style='List Bullet')
    doc.add_heading("Conflict of Interest", level=2)
    for c in minutes_data.get("conflict_of_interest", []):
        doc.add_paragraph(f"• {c}", style='List Bullet')
    path = os.path.join(TEMP_DIR, filename)
    doc.save(path)
    return {"message": "Word exported", "filename": filename}


# ---------- 6. 导出 JSON ----------
@app.post("/export_json")
async def export_json(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.json"
    path = os.path.join(TEMP_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(minutes_data, f, ensure_ascii=False, indent=2)
    return {"message": "JSON exported", "filename": filename}


# ---------- 7. 导出 PDF ----------
@app.post("/export_pdf")
async def export_pdf(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.pdf"
    path = os.path.join(TEMP_DIR, filename)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, f"Meeting Summary\n\nSummary:\n{minutes_data.get('summary', '')}\n\n")

    def add_section(title, items):
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(0, 10, title, ln=True)
        pdf.set_font("Arial", size=12)
        for item in items:
            pdf.multi_cell(0, 8, f"- {item}")
        pdf.ln(2)

    add_section("Agenda Items", minutes_data.get("agenda_items", []))
    add_section("Actions", minutes_data.get("action", []))
    add_section("Decisions", minutes_data.get("decision", []))
    add_section("Conflict of Interest", minutes_data.get("conflict_of_interest", []))

    pdf.output(path)
    return {"message": "PDF exported", "filename": filename}


# ---------- 8. 下载文件 ----------
@app.get("/download/{filename}")
async def download_file(filename: str):
    path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(path):
        return {"error": "文件不存在"}
    media_type, _ = mimetypes.guess_type(path)
    return FileResponse(path, media_type=media_type or "application/octet-stream", filename=filename)


# ---------- ✅ 9. 语义搜索接口 ----------
@app.post("/semantic_search")
def semantic_search(request: SearchRequest):
    try:
        results = hs.search(request.query)
        return results
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
