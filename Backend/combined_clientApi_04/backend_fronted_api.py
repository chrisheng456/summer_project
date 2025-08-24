from fastapi import FastAPI, UploadFile, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
import shutil
import datetime
import whisper
from docx import Document

app = FastAPI()

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins during frontend debugging
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary folder to save files
TEMP_DIR = "temp"
os.makedirs(TEMP_DIR, exist_ok=True)

# Load Whisper model
model = whisper.load_model("base")


# 1. Upload audio file
@app.post("/upload_audio")
async def upload_audio(file: UploadFile):
    save_path = os.path.join(TEMP_DIR, file.filename)
    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    return {"message": "File uploaded successfully", "filename": file.filename}


# 2. Transcribe audio using Whisper
@app.post("/transcribe")
async def transcribe_audio(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)
    try:
        result = model.transcribe(file_path)
        return {"transcription": result["text"]}
    except Exception as e:
        return {"error": str(e)}


# 3. Export meeting summary as Markdown
@app.post("/export_markdown")
async def export_markdown(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.md"

    md = f"""# 📝 Meeting Summary

## Summary  
{minutes_data.get('summary', '')}

## Key Points  
""" + "\n".join(f"- {kp}" for kp in minutes_data.get("key_points", [])) + """

## Action Items  
""" + "\n".join(f"- [ ] {ai}" for ai in minutes_data.get("action_items", [])) + f"""

## Sentiment  
{minutes_data.get("sentiment", "N/A")}
"""

    save_path = os.path.join(TEMP_DIR, filename)
    with open(save_path, "w", encoding="utf-8") as f:
        f.write(md)

    return {"message": "Markdown exported", "filename": filename}


# 4. Export meeting summary as Word document
@app.post("/export_word")
async def export_word(minutes_data: dict = Body(...)):
    now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"meeting_summary_{now}.docx"
    doc = Document()

    doc.add_heading("Meeting Summary", level=1)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(minutes_data.get("summary", ""))

    doc.add_heading("Key Points", level=2)
    for item in minutes_data.get("key_points", []):
        doc.add_paragraph(f"• {item}", style='List Bullet')

    doc.add_heading("Action Items", level=2)
    for item in minutes_data.get("action_items", []):
        doc.add_paragraph(f"☐ {item}", style='List Bullet')

    doc.add_heading("Sentiment", level=2)
    doc.add_paragraph(minutes_data.get("sentiment", "N/A"))

    save_path = os.path.join(TEMP_DIR, filename)
    doc.save(save_path)

    return {"message": "Word exported", "filename": filename}


# 5. Download exported files (Markdown / Word)
@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)
    if not os.path.exists(file_path):
        return {"error": "File does not exist"}
    return FileResponse(file_path, media_type="application/octet-stream", filename=filename)
